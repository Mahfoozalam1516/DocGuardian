import streamlit as st
import os
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain_community.llms import HuggingFaceHub
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from PyPDF2 import PdfReader
import docx
import pandas as pd
import time
import json
from datetime import datetime
import plotly.express as px
from pathlib import Path
import base64

# Configure Streamlit page
st.set_page_config(
    page_title="📚 Enhanced Document Q&A Chatbot",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
        display: flex;
    }
    .chat-message.user {
        background-color: #e6f3ff;
    }
    .chat-message.bot {
        background-color: #f0f2f6;
    }
    .chat-icon {
        width: 50px;
        margin-right: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state variables
def init_session_state():
    session_vars = {
        'conversation': None,
        'chat_history': [],
        'token_validated': False,
        'processed_files': [],
        'total_tokens': 0,
        'conversation_start_time': None,
        'chat_stats': {'total_questions': 0, 'avg_response_time': 0},
        'document_stats': {'total_docs': 0, 'total_pages': 0},
        'api_key': None
    }
    
    for var, value in session_vars.items():
        if var not in st.session_state:
            st.session_state[var] = value

def get_file_text(file):
    """Extract text from various file formats"""
    text = ""
    file_extension = Path(file.name).suffix.lower()
    
    try:
        if file_extension == '.pdf':
            reader = PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
            st.session_state.document_stats['total_pages'] += len(reader.pages)
            
        elif file_extension == '.docx':
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + '\n'
            st.session_state.document_stats['total_pages'] += len(doc.paragraphs) // 40  # Approximate
            
        elif file_extension == '.txt':
            text = file.getvalue().decode('utf-8')
            st.session_state.document_stats['total_pages'] += len(text.split('\n')) // 40  # Approximate
            
        return text
    except Exception as e:
        st.error(f"Error processing {file.name}: {str(e)}")
        return ""

def get_text_chunks(text):
    """Split text into chunks"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    chunks = text_splitter.split_text(text)
    return chunks

def validate_api_key(api_key):
    """Validate HuggingFace API token"""
    try:
        os.environ['HUGGINGFACEHUB_API_TOKEN'] = api_key
        HuggingFaceHub(
            repo_id="google/flan-t5-base",
            model_kwargs={"temperature": 0.5, "max_length": 512},
            huggingfacehub_api_token=api_key
        )
        return True
    except Exception as e:
        return False

def get_conversation_chain(vectorstore):
    """Initialize conversation chain"""
    llm = HuggingFaceHub(
        repo_id="google/flan-t5-base",
        model_kwargs={
            "temperature": 0.7,
            "max_length": 512,
            "top_p": 0.9
        },
        huggingfacehub_api_token=st.session_state.api_key
    )
    
    memory = ConversationBufferMemory(
        memory_key='chat_history',
        return_messages=True,
        output_key='answer'
    )
    
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
        memory=memory,
        return_source_documents=True,
        verbose=True
    )
    
    return conversation_chain

def handle_user_input(question):
    """Process user input and generate response"""
    if st.session_state.conversation is None:
        st.warning("Please upload documents before asking questions.")
        return
    
    start_time = time.time()
    
    try:
        response = st.session_state.conversation({
            'question': question,
            'chat_history': st.session_state.chat_history
        })
        
        # Update statistics
        response_time = time.time() - start_time
        st.session_state.chat_stats['total_questions'] += 1
        st.session_state.chat_stats['avg_response_time'] = (
            (st.session_state.chat_stats['avg_response_time'] * 
             (st.session_state.chat_stats['total_questions'] - 1) + 
             response_time) / st.session_state.chat_stats['total_questions']
        )
        
        # Display chat message
        with st.chat_message("user"):
            st.write(question)
        
        with st.chat_message("assistant"):
            st.write(response['answer'])
            
            # Display source documents
            if 'source_documents' in response and response['source_documents']:
                with st.expander("View Sources"):
                    for i, doc in enumerate(response['source_documents']):
                        st.markdown(f"**Source {i+1}:**")
                        st.markdown(doc.page_content)
        
        # Update chat history
        st.session_state.chat_history.append((question, response['answer']))
        
    except Exception as e:
        st.error(f"Error generating response: {str(e)}")

def display_chat_history():
    """Display chat history with download option"""
    if st.session_state.chat_history:
        st.subheader("Chat History")
        
        # Convert chat history to DataFrame
        df = pd.DataFrame(st.session_state.chat_history, columns=['Question', 'Answer'])
        
        # Download button
        csv = df.to_csv(index=False)
        b64 = base64.b64encode(csv.encode()).decode()
        href = f'<a href="data:file/csv;base64,{b64}" download="chat_history.csv">Download Chat History</a>'
        st.markdown(href, unsafe_allow_html=True)
        
        # Display history in expander
        with st.expander("View Chat History"):
            st.dataframe(df)

def display_stats():
    """Display chat and document statistics"""
    if st.session_state.chat_stats['total_questions'] > 0:
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Total Questions", st.session_state.chat_stats['total_questions'])
        
        with col2:
            st.metric("Avg Response Time", f"{st.session_state.chat_stats['avg_response_time']:.2f}s")
            
        with col3:
            st.metric("Documents Processed", st.session_state.document_stats['total_docs'])
        
        # Create usage graph
        times = [i for i in range(st.session_state.chat_stats['total_questions'])]
        questions = [1 for _ in range(st.session_state.chat_stats['total_questions'])]
        
        fig = px.line(
            x=times,
            y=questions,
            title="Questions Over Time",
            labels={'x': 'Question Number', 'y': 'Questions'}
        )
        st.plotly_chart(fig)

def main():
    init_session_state()
    st.title("📚 Enhanced Document Q&A Chatbot")
    
    # Sidebar
    with st.sidebar:
        st.header("Settings & Info")
        st.markdown("---")
        
        # API Key Input
        if not st.session_state.token_validated:
            st.subheader("🔑 API Configuration")
            api_key = st.text_input("Enter HuggingFace API Key:", type="password")
            if st.button("Validate API Key"):
                if api_key and validate_api_key(api_key):
                    st.session_state.api_key = api_key
                    st.session_state.token_validated = True
                    st.success("✅ API Key validated successfully!")
                else:
                    st.error("❌ Invalid API Key")
        
        # File upload (only show if API key is validated)
        if st.session_state.token_validated:
            uploaded_files = st.file_uploader(
                "Upload Documents",
                type=['pdf', 'docx', 'txt'],
                accept_multiple_files=True,
                help="Supported formats: PDF, DOCX, TXT"
            )
            
            if uploaded_files:
                process_button = st.button("Process Documents")
                if process_button:
                    with st.spinner("Processing documents..."):
                        # Process each file
                        all_text = ""
                        for file in uploaded_files:
                            if file.name not in st.session_state.processed_files:
                                text = get_file_text(file)
                                if text:
                                    all_text += text + "\n\n"
                                    st.session_state.processed_files.append(file.name)
                                    st.session_state.document_stats['total_docs'] += 1
                        
                        if all_text:
                            # Create text chunks
                            text_chunks = get_text_chunks(all_text)
                            
                            # Create embeddings
                            embeddings = HuggingFaceEmbeddings(
                                model_name="sentence-transformers/all-MiniLM-L6-v2"
                            )
                            
                            # Create vector store
                            vectorstore = FAISS.from_texts(texts=text_chunks, embedding=embeddings)
                            
                            # Create conversation chain
                            st.session_state.conversation = get_conversation_chain(vectorstore)
                            
                            st.success("Documents processed successfully!")
                            
                            # Initialize conversation start time
                            if not st.session_state.conversation_start_time:
                                st.session_state.conversation_start_time = datetime.now()
    
    # Main chat interface
    if st.session_state.token_validated:
        if st.session_state.conversation:
            # Display statistics
            display_stats()
            
            # Chat interface
            st.subheader("Ask Questions")
            user_question = st.text_input("Type your question here:")
            
            if user_question:
                handle_user_input(user_question)
            
            # Display chat history
            display_chat_history()
        
        # Display processed files
        if st.session_state.processed_files:
            with st.expander("Processed Files"):
                for file in st.session_state.processed_files:
                    st.write(f"✅ {file}")
    else:
        st.info("Please enter your HuggingFace API Key in the sidebar to get started.")

if __name__ == '__main__':
    main()